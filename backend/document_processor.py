import PyPDF2
from docx import Document
import io
import re

class DocumentProcessor:
    """Handle different document formats and text extraction"""
    
    SUPPORTED_FORMATS = ['.txt', '.pdf', '.docx', '.md']
    
    def extract_text(self, file):
        """Extract text from uploaded file"""
        filename = file.filename.lower()
        
        if filename.endswith('.txt') or filename.endswith('.md'):
            return self._extract_from_txt(file)
        elif filename.endswith('.pdf'):
            return self._extract_from_pdf(file)
        elif filename.endswith('.docx'):
            return self._extract_from_docx(file)
        else:
            raise ValueError(f"Unsupported file format. Supported formats: {', '.join(self.SUPPORTED_FORMATS)}")
    
    def _extract_from_txt(self, file):
        """Extract text from TXT or Markdown file"""
        try:
            content = file.read().decode('utf-8')
            return self._clean_text(content)
        except UnicodeDecodeError:
            # Try different encoding
            file.seek(0)
            try:
                content = file.read().decode('latin-1')
                return self._clean_text(content)
            except Exception as e:
                raise ValueError(f"Could not decode text file: {str(e)}")
    
    def _extract_from_pdf(self, file):
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            
            if len(pdf_reader.pages) == 0:
                raise ValueError("PDF file is empty")
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return self._clean_text(text)
            
        except Exception as e:
            raise ValueError(f"Could not extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file):
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file.read()))
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text found in DOCX file")
            
            return self._clean_text(text)
            
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters but keep letters, numbers, punctuation
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def get_word_count(self, text):
        """Get accurate word count"""
        if not text:
            return 0
        return len(text.split())
    
    def get_text_preview(self, text, length=200):
        """Get preview of text"""
        if not text:
            return ""
        if len(text) <= length:
            return text
        return text[:length] + "..."